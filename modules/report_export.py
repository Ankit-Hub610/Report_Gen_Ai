"""
report_export.py
-----------------
Full Analysis (Page 1 + Page 2) export to PDF / Word (.docx) / Excel (.xlsx).

Design goal: reuse the EXACT same computed dicts the "Full Analysis" screen
already shows (`facts` from intel_engine.build_facts_bundle, `ir` from
intel_engine.generate_insights_and_recommendations, plus roles/cleaning_log/
derived_log) — never recompute or invent a number here. If a number changes
on screen, the same change shows up in every exported file, automatically.

Public functions:
    build_full_analysis_pdf(facts, ir, health, roles, cleaning_log, derived_log,
                             theme, which="both") -> bytes
    build_full_analysis_docx(facts, ir, health, roles, cleaning_log, derived_log,
                              which="both") -> bytes
    build_full_analysis_xlsx(facts, ir, health, roles, cleaning_log, derived_log,
                              which="both") -> bytes

`which` is one of "page1", "page2", "both" — mirrors the on-screen page
radio, so a user who only wants Page 2 (Summary & Recommendations) doesn't
have to download Page 1's detail too.
"""

import io
from datetime import datetime

import pandas as pd

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                                 PageBreak, ListFlowable, ListItem)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter


# ----------------------------------------------------------------------------
# Shared helpers — pull the same numbers the screen shows, in one place, so
# PDF/DOCX/XLSX can't ever drift from each other.
# ----------------------------------------------------------------------------
def _fmt(v):
    if v is None:
        return "\u2014"
    try:
        if isinstance(v, float) and v == int(v):
            v = int(v)
        return f"{v:,}" if isinstance(v, int) else f"{v:,.2f}"
    except Exception:
        return str(v)


def _kpi_band(facts, ir_mode="page1"):
    f = facts["financials"]
    if ir_mode == "page1":
        return [
            ("Total Revenue", _fmt(f.get("total_revenue"))),
            ("Total Orders", _fmt(f.get("total_orders"))),
            ("Customers", _fmt(f.get("customer_count"))),
            ("Avg Order Value", _fmt(f.get("avg_order_value"))),
        ]
    return [
        ("Total Revenue", _fmt(f.get("total_revenue"))),
        ("Total Orders", _fmt(f.get("total_orders"))),
        ("Customers", _fmt(f.get("customer_count"))),
        ("Profit Margin", f"{f.get('profit_margin_pct')}%" if f.get("profit_margin_pct") is not None else "N/A"),
    ]


def _strip_md_bold(text):
    """Screen text uses **bold** markdown; PDF/DOCX render real bold instead,
    so strip the literal asterisks for the plain-text runs and return the
    plain string (bold runs are handled separately where it matters)."""
    return (text or "").replace("**", "")


BREAKDOWN_LABELS = {"product": "Product", "customer": "Customer", "location": "Location", "channel": "Channel"}


# ============================================================================
# PDF
# ============================================================================
def build_full_analysis_pdf(facts, ir, health, roles, cleaning_log, derived_log, theme, which="both"):
    buf = io.BytesIO()
    accent_hex = theme.get("accent_color", "#2C6E49")
    font_hex = theme.get("font_color", "#111111")
    font_name = theme.get("font_name", "Helvetica")
    accent = colors.HexColor(accent_hex)
    font_color = colors.HexColor(font_hex)
    grid = colors.HexColor("#CCCCCC")

    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.6 * cm, rightMargin=1.6 * cm,
                             topMargin=1.4 * cm, bottomMargin=1.4 * cm,
                             title="Full Analysis Report")
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("T", parent=styles["Title"], textColor=accent, fontName=font_name, fontSize=20)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=accent, fontName=font_name, fontSize=13,
                         spaceBefore=12, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["Normal"], textColor=font_color, fontName=font_name, fontSize=9.5,
                           leading=13)
    small = ParagraphStyle("Small", parent=styles["Normal"], textColor=font_color, fontName=font_name, fontSize=8.5,
                            leading=11)

    def kpi_table(rows):
        data = [[Paragraph(f"<b>{lbl}</b>", small) for lbl, _ in rows],
                [Paragraph(val, body) for _, val in rows]]
        t = Table(data, colWidths=[doc.width / len(rows)] * len(rows))
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.6, grid), ("INNERGRID", (0, 0), (-1, -1), 0.4, grid),
            ("BACKGROUND", (0, 0), (-1, 0), accent),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        return t

    def df_table(rows, header):
        data = [header] + rows
        t = Table(data, repeatRows=1, colWidths=[doc.width / len(header)] * len(header))
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.5, grid), ("INNERGRID", (0, 0), (-1, -1), 0.3, grid),
            ("BACKGROUND", (0, 0), (-1, 0), accent), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8), ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    story = [Paragraph("Full Analysis Report", title_style),
             Paragraph(f"Generated {datetime.now().strftime('%d %b %Y, %H:%M')}", small),
             Spacer(1, 10)]

    # ---------------- PAGE 1 ----------------
    if which in ("page1", "both"):
        story.append(Paragraph(f"Business Health: {health['label']}", h2))
        if health.get("reasons"):
            story.append(Paragraph(" \u2022 ".join(health["reasons"]), small))
        story.append(Spacer(1, 6))
        story.append(kpi_table(_kpi_band(facts, "page1")))

        story.append(Paragraph("Data Understanding & Cleaning", h2))
        for line in cleaning_log:
            story.append(Paragraph(f"\u2022 {_strip_md_bold(line)}", small))
        story.append(Paragraph(f"<b>Data Quality Score: {facts['quality']['score']}/100</b>", body))
        for issue in facts["quality"]["issues"]:
            story.append(Paragraph(f"\u2022 {_strip_md_bold(issue)}", small))

        story.append(Paragraph("Calculated / Added Columns", h2))
        for line in derived_log:
            story.append(Paragraph(f"\u2022 {_strip_md_bold(line)}", small))
        detected = ", ".join(f"{k}={v}" for k, v in roles.items() if v) or "\u2014"
        story.append(Paragraph(f"Detected column roles: {detected}", small))

        t, fc = facts["trend"], facts["forecast"]
        story.append(Paragraph("Revenue Trend & Forecast", h2))
        if t.get("available"):
            story.append(Paragraph(
                f"Best period: <b>{t['best_period']}</b> ({_fmt(t['best_period_value'])}) &nbsp;|&nbsp; "
                f"Worst period: <b>{t['worst_period']}</b> ({_fmt(t['worst_period_value'])})", body))
            if t.get("overall_change_pct") is not None:
                story.append(Paragraph(
                    f"Overall change: <b>{t['overall_change_pct']:+.1f}%</b>"
                    + (f" &nbsp;|&nbsp; CAGR: <b>{t.get('cagr_pct')}%</b>" if t.get("cagr_pct") is not None else ""),
                    body))
            rows = [[p, _fmt(v), (f"{m:+.1f}%" if m is not None else "\u2014")]
                    for p, v, m in zip(t["periods"], t["values"], t["mom_growth_pct"])]
            story.append(Spacer(1, 4))
            story.append(df_table(rows, ["Period", "Revenue", "MoM Growth %"]))
        else:
            story.append(Paragraph(_strip_md_bold(t.get("reason", "Trend not available.")), body))

        if fc.get("available"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                f"Forecast (next {len(fc['forecast_periods'])} months) \u2014 method: {fc['method']}, "
                f"confidence: <b>{fc['confidence']}</b> (R\u00b2={fc['r2']}), direction: <b>{fc['direction']}</b>", body))
            rows = [[p, _fmt(v), typ] for p, v, typ in zip(
                fc["history_periods"] + fc["forecast_periods"],
                fc["history_values"] + fc["forecast_values"],
                ["Actual"] * len(fc["history_periods"]) + ["Forecast"] * len(fc["forecast_periods"]))]
            story.append(df_table(rows, ["Period", "Value", "Type"]))
            story.append(Paragraph("Rows marked Forecast are estimates, not actual results.", small))

        if facts["anomalies"]:
            story.append(Paragraph("Anomalies Detected", h2))
            cols = list(facts["anomalies"][0].keys())
            rows = [[str(a.get(c, "")) for c in cols] for a in facts["anomalies"]]
            story.append(df_table(rows, cols))

        if facts["correlations"]:
            story.append(Paragraph("Correlations", h2))
            cols = list(facts["correlations"][0].keys())
            rows = [[str(c.get(cc, "")) for cc in cols] for c in facts["correlations"]]
            story.append(df_table(rows, cols))
            story.append(Paragraph("Correlation does not prove causation.", small))

        for key, label in BREAKDOWN_LABELS.items():
            b = facts["breakdowns"].get(key, {})
            if not b.get("available"):
                continue
            story.append(Paragraph(f"Top / Bottom by {label} ({b['dimension']})", h2))
            story.append(Paragraph(
                f"Measured on <b>{b['measure']}</b> \u2022 {b['unique_count']} unique values \u2022 "
                f"top-5 share of total: <b>{b['top5_share_pct']}%</b>", small))
            top_cols = list(b["top"][0].keys()) if b["top"] else []
            if top_cols:
                story.append(Paragraph("Top performers", body))
                rows = [[str(r.get(c, "")) for c in top_cols] for r in b["top"]]
                story.append(df_table(rows, top_cols))
            if b.get("bottom"):
                story.append(Paragraph("Bottom performers", body))
                rows = [[str(r.get(c, "")) for c in top_cols] for r in b["bottom"]]
                story.append(df_table(rows, top_cols))

    # ---------------- PAGE 2 ----------------
    if which == "both":
        story.append(PageBreak())
    if which in ("page2", "both"):
        story.append(Paragraph(f"Summary & Recommendations \u2014 Business Health: {health['label']}", h2))
        story.append(kpi_table(_kpi_band(facts, "page2")))

        story.append(Paragraph("Past Performance", h2))
        story.append(Paragraph(_strip_md_bold(ir["past_summary"]), body))
        story.append(Paragraph("Future Outlook", h2))
        story.append(Paragraph(_strip_md_bold(ir["future_summary"]), body))

        story.append(Paragraph("Key Insights", h2))
        story.append(ListFlowable(
            [ListItem(Paragraph(_strip_md_bold(x), body)) for x in ir["key_insights"]],
            bulletType="bullet", start="\u2022"))

        story.append(Paragraph("Recommended Actions", h2))
        story.append(ListFlowable(
            [ListItem(Paragraph(_strip_md_bold(x), body)) for x in ir["recommended_actions"]],
            bulletType="1"))

    doc.build(story)
    return buf.getvalue()


# ============================================================================
# WORD (.docx)
# ============================================================================
def build_full_analysis_docx(facts, ir, health, roles, cleaning_log, derived_log, which="both"):
    d = Document()

    def h1(text):
        p = d.add_heading(text, level=1)
        return p

    def h2(text):
        d.add_heading(text, level=2)

    def para(text, bold=False, size=10):
        p = d.add_paragraph()
        r = p.add_run(_strip_md_bold(text))
        r.bold = bold
        r.font.size = Pt(size)
        return p

    def bullets(items):
        for it in items:
            p = d.add_paragraph(_strip_md_bold(it), style="List Bullet")

    def numbered(items):
        for it in items:
            d.add_paragraph(_strip_md_bold(it), style="List Number")

    def kpi_table(rows):
        t = d.add_table(rows=2, cols=len(rows))
        t.style = "Light Grid Accent 1"
        for i, (lbl, val) in enumerate(rows):
            t.rows[0].cells[i].text = lbl
            t.rows[1].cells[i].text = val

    def df_table(rows, header):
        t = d.add_table(rows=1, cols=len(header))
        t.style = "Light Grid Accent 1"
        for i, c in enumerate(header):
            t.rows[0].cells[i].text = str(c)
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = str(v)

    title = d.add_heading("Full Analysis Report", level=0)
    para(f"Generated {datetime.now().strftime('%d %b %Y, %H:%M')}", size=9)

    if which in ("page1", "both"):
        h1(f"Business Health: {health['label']}")
        if health.get("reasons"):
            para(" \u2022 ".join(health["reasons"]), size=9)
        kpi_table(_kpi_band(facts, "page1"))

        h2("Data Understanding & Cleaning")
        bullets(cleaning_log)
        para(f"Data Quality Score: {facts['quality']['score']}/100", bold=True)
        bullets(facts["quality"]["issues"])

        h2("Calculated / Added Columns")
        bullets(derived_log)
        detected = ", ".join(f"{k}={v}" for k, v in roles.items() if v) or "\u2014"
        para(f"Detected column roles: {detected}", size=9)

        t, fc = facts["trend"], facts["forecast"]
        h2("Revenue Trend & Forecast")
        if t.get("available"):
            para(f"Best period: {t['best_period']} ({_fmt(t['best_period_value'])})  |  "
                 f"Worst period: {t['worst_period']} ({_fmt(t['worst_period_value'])})")
            if t.get("overall_change_pct") is not None:
                para(f"Overall change: {t['overall_change_pct']:+.1f}%"
                     + (f"  |  CAGR: {t.get('cagr_pct')}%" if t.get("cagr_pct") is not None else ""))
            rows = [[p, _fmt(v), (f"{m:+.1f}%" if m is not None else "\u2014")]
                    for p, v, m in zip(t["periods"], t["values"], t["mom_growth_pct"])]
            df_table(rows, ["Period", "Revenue", "MoM Growth %"])
        else:
            para(t.get("reason", "Trend not available."))

        if fc.get("available"):
            para(f"Forecast (next {len(fc['forecast_periods'])} months) \u2014 method: {fc['method']}, "
                 f"confidence: {fc['confidence']} (R2={fc['r2']}), direction: {fc['direction']}")
            rows = [[p, _fmt(v), typ] for p, v, typ in zip(
                fc["history_periods"] + fc["forecast_periods"],
                fc["history_values"] + fc["forecast_values"],
                ["Actual"] * len(fc["history_periods"]) + ["Forecast"] * len(fc["forecast_periods"]))]
            df_table(rows, ["Period", "Value", "Type"])
            para("Rows marked Forecast are estimates, not actual results.", size=9)

        if facts["anomalies"]:
            h2("Anomalies Detected")
            cols = list(facts["anomalies"][0].keys())
            rows = [[str(a.get(c, "")) for c in cols] for a in facts["anomalies"]]
            df_table(rows, cols)

        if facts["correlations"]:
            h2("Correlations")
            cols = list(facts["correlations"][0].keys())
            rows = [[str(c.get(cc, "")) for cc in cols] for c in facts["correlations"]]
            df_table(rows, cols)
            para("Correlation does not prove causation.", size=9)

        for key, label in BREAKDOWN_LABELS.items():
            b = facts["breakdowns"].get(key, {})
            if not b.get("available"):
                continue
            h2(f"Top / Bottom by {label} ({b['dimension']})")
            para(f"Measured on {b['measure']} \u2022 {b['unique_count']} unique values \u2022 "
                 f"top-5 share of total: {b['top5_share_pct']}%", size=9)
            top_cols = list(b["top"][0].keys()) if b["top"] else []
            if top_cols:
                para("Top performers", bold=True)
                rows = [[str(r.get(c, "")) for c in top_cols] for r in b["top"]]
                df_table(rows, top_cols)
            if b.get("bottom"):
                para("Bottom performers", bold=True)
                rows = [[str(r.get(c, "")) for c in top_cols] for r in b["bottom"]]
                df_table(rows, top_cols)

    if which == "both":
        d.add_page_break()
    if which in ("page2", "both"):
        h1(f"Summary & Recommendations \u2014 Business Health: {health['label']}")
        kpi_table(_kpi_band(facts, "page2"))

        h2("Past Performance")
        para(ir["past_summary"])
        h2("Future Outlook")
        para(ir["future_summary"])
        h2("Key Insights")
        bullets(ir["key_insights"])
        h2("Recommended Actions")
        numbered(ir["recommended_actions"])

    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


# ============================================================================
# EXCEL (.xlsx)
# ============================================================================
def build_full_analysis_xlsx(facts, ir, health, roles, cleaning_log, derived_log, which="both"):
    wb = Workbook()
    header_fill = PatternFill("solid", fgColor="2C6E49")
    header_font = Font(bold=True, color="FFFFFF")
    bold = Font(bold=True)

    def new_sheet(name):
        ws = wb.create_sheet(title=name[:31])
        return ws

    def write_kv(ws, row, key, val):
        ws.cell(row=row, column=1, value=key).font = bold
        ws.cell(row=row, column=2, value=val)
        return row + 1

    def write_table(ws, start_row, header, rows, title=None):
        r = start_row
        if title:
            ws.cell(row=r, column=1, value=title).font = Font(bold=True, size=12)
            r += 1
        for c, h in enumerate(header, 1):
            cell = ws.cell(row=r, column=c, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")
        r += 1
        for row_data in rows:
            for c, v in enumerate(row_data, 1):
                ws.cell(row=r, column=c, value=v)
            r += 1
        for c in range(1, len(header) + 1):
            ws.column_dimensions[get_column_letter(c)].width = 18
        return r + 1

    ws0 = wb.active
    ws0.title = "Summary"
    ws0.cell(row=1, column=1, value="Full Analysis Report").font = Font(bold=True, size=16)
    r = 3
    r = write_kv(ws0, r, "Generated", datetime.now().strftime("%d %b %Y, %H:%M"))
    r = write_kv(ws0, r, "Business Health", health["label"])
    for lbl, val in _kpi_band(facts, "page1"):
        r = write_kv(ws0, r, lbl, val)
    ws0.column_dimensions["A"].width = 22
    ws0.column_dimensions["B"].width = 30

    if which in ("page1", "both"):
        ws1 = new_sheet("Data Quality & Cleaning")
        r = 1
        ws1.cell(row=r, column=1, value=f"Data Quality Score: {facts['quality']['score']}/100").font = bold
        r += 2
        r = write_table(ws1, r, ["Cleaning notes"], [[l] for l in cleaning_log])
        r = write_table(ws1, r, ["Quality issues"], [[i] for i in facts["quality"]["issues"]])
        r = write_table(ws1, r, ["Calculated / added columns"], [[l] for l in derived_log])
        detected = [[f"{k}", f"{v}"] for k, v in roles.items() if v]
        write_table(ws1, r, ["Role", "Column"], detected, title="Detected column roles")

        t, fc = facts["trend"], facts["forecast"]
        if t.get("available"):
            ws2 = new_sheet("Trend")
            rows = [[p, v, m] for p, v, m in zip(t["periods"], t["values"], t["mom_growth_pct"])]
            write_table(ws2, 1, ["Period", "Revenue", "MoM Growth %"], rows,
                        title=f"Best: {t['best_period']} | Worst: {t['worst_period']}")
        if fc.get("available"):
            ws3 = new_sheet("Forecast")
            rows = [[p, v, typ] for p, v, typ in zip(
                fc["history_periods"] + fc["forecast_periods"],
                fc["history_values"] + fc["forecast_values"],
                ["Actual"] * len(fc["history_periods"]) + ["Forecast"] * len(fc["forecast_periods"]))]
            write_table(ws3, 1, ["Period", "Value", "Type"], rows,
                        title=f"Method: {fc['method']} | Confidence: {fc['confidence']} | Direction: {fc['direction']}")

        if facts["anomalies"]:
            ws4 = new_sheet("Anomalies")
            cols = list(facts["anomalies"][0].keys())
            rows = [[a.get(c, "") for c in cols] for a in facts["anomalies"]]
            write_table(ws4, 1, cols, rows)

        if facts["correlations"]:
            ws5 = new_sheet("Correlations")
            cols = list(facts["correlations"][0].keys())
            rows = [[c.get(cc, "") for cc in cols] for c in facts["correlations"]]
            write_table(ws5, 1, cols, rows)

        for key, label in BREAKDOWN_LABELS.items():
            b = facts["breakdowns"].get(key, {})
            if not b.get("available"):
                continue
            wsb = new_sheet(f"{label} Breakdown")
            top_cols = list(b["top"][0].keys()) if b["top"] else []
            r = write_table(wsb, 1, top_cols, [[r_.get(c, "") for c in top_cols] for r_ in b["top"]],
                             title=f"Top {label} (measured on {b['measure']})")
            if b.get("bottom"):
                write_table(wsb, r, top_cols, [[r_.get(c, "") for c in top_cols] for r_ in b["bottom"]],
                            title=f"Bottom {label}")

    if which in ("page2", "both"):
        ws6 = new_sheet("Summary & Recommendations")
        r = 1
        for lbl, val in _kpi_band(facts, "page2"):
            r = write_kv(ws6, r, lbl, val)
        r += 1
        ws6.cell(row=r, column=1, value="Past Performance").font = Font(bold=True, size=12)
        r += 1
        ws6.cell(row=r, column=1, value=_strip_md_bold(ir["past_summary"]))
        r += 2
        ws6.cell(row=r, column=1, value="Future Outlook").font = Font(bold=True, size=12)
        r += 1
        ws6.cell(row=r, column=1, value=_strip_md_bold(ir["future_summary"]))
        r += 2
        r = write_table(ws6, r, ["Key Insights"], [[_strip_md_bold(x)] for x in ir["key_insights"]])
        write_table(ws6, r, ["#", "Recommended Action"],
                    [[i, _strip_md_bold(x)] for i, x in enumerate(ir["recommended_actions"], 1)])
        ws6.column_dimensions["A"].width = 90

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()
